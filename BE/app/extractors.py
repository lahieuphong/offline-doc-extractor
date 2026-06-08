import os
from pathlib import Path
from typing import List, Literal, Optional, Tuple

import fitz  # PyMuPDF
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg"}
OCR_TIMEOUT_SEC = int(os.getenv("OCR_TIMEOUT_SEC", "45"))
OCR_MIN_TEXT_LEN = max(0, int(os.getenv("OCR_MIN_TEXT_LEN", "120")))
OCR_MAX_PAGES_PER_PDF = max(0, int(os.getenv("OCR_MAX_PAGES_PER_PDF", "4")))
OCR_RENDER_ZOOM = float(os.getenv("OCR_RENDER_ZOOM", "2.5"))
OCR_TESSERACT_PSM = os.getenv("OCR_TESSERACT_PSM", "6")
OCR_TARGETED_REGIONS = os.getenv("OCR_TARGETED_REGIONS", "1").lower() not in {"0", "false", "no"}
OCR_TARGETED_ZOOM = float(os.getenv("OCR_TARGETED_ZOOM", "3.5"))
OCR_TARGETED_TIMEOUT_SEC = int(os.getenv("OCR_TARGETED_TIMEOUT_SEC", "12"))
OCR_MULTI_VARIANT = os.getenv("OCR_MULTI_VARIANT", "1").lower() not in {"0", "false", "no"}
OCR_TESSERACT_OEM = os.getenv("OCR_TESSERACT_OEM", "1")
OCR_LANG = os.getenv("OCR_LANG", "vie+eng")
OCR_FULL_PAGE = os.getenv("OCR_FULL_PAGE", "0").lower() not in {"0", "false", "no"}
OCR_TARGETED_ALWAYS = os.getenv("OCR_TARGETED_ALWAYS", "1").lower() not in {"0", "false", "no"}

TEXT_LAYER_NOISE_MARKERS = (
    "người ký:",
    "email:",
    "cơ quan:",
    "thời gian ký:",
    "cổng thông tin điện tử chính phủ",
)


def is_text_layer_noise(text_layer: str) -> bool:
    normalized = " ".join((text_layer or "").lower().split())
    if not normalized:
        return True
    marker_hits = sum(1 for marker in TEXT_LAYER_NOISE_MARKERS if marker in normalized)
    if marker_hits >= 2 and len(normalized) < 500:
        return True
    if len(normalized) < OCR_MIN_TEXT_LEN:
        return True
    return False


def _ocr_config(psm: str) -> str:
    # preserve_interword_spaces helps keep document codes such as 971/QĐ-TTg.
    return f"--oem {OCR_TESSERACT_OEM} --psm {psm} -c preserve_interword_spaces=1"


def _preprocess_variants(image: Image.Image, *, aggressive: bool = False) -> List[Image.Image]:
    """Create a small set of OCR variants and pick the best output later.

    Government scans often contain stamps/seals.  One image variant rarely wins
    for every file, so we try a grayscale/contrast version and a thresholded
    version.  No external online service is used.
    """
    rgb = image.convert("RGB")
    gray = ImageOps.grayscale(rgb)
    gray = ImageOps.autocontrast(gray)
    contrast = ImageEnhance.Contrast(gray).enhance(2.2 if aggressive else 1.7)
    sharp = contrast.filter(ImageFilter.SHARPEN)
    if not OCR_MULTI_VARIANT:
        return [sharp]
    variants = [rgb, sharp]
    threshold_value = 178 if aggressive else 188
    bw = sharp.point(lambda p: 255 if p > threshold_value else 0)
    variants.append(bw)
    return variants


def _score_ocr_text(text: str) -> float:
    clean = text or ""
    letters = len([ch for ch in clean if ch.isalpha()])
    vi_chars = len([ch for ch in clean if ch in "ăâđêôơưĂÂĐÊÔƠƯáàảãạéèẻẽẹíìỉĩịóòỏõọúùủũụýỳỷỹỵếềểễệốồổỗộớờởỡợứừửữựắằẳẵặ"])
    digits = len([ch for ch in clean if ch.isdigit()])
    bad = len([ch for ch in clean if ch in "<>|\\{}[]`^~§¤�*_="])
    upper = clean.upper()
    noise_markers = sum(marker in upper for marker in ["OCR_TIMEOUT", "GONG THONG", "CHINA PH", "DIENT", "TNĐIỆN"])
    return letters + vi_chars * 2 + min(digits, 20) - bad * 8 - noise_markers * 120


def _image_to_string_best(image: Image.Image, *, psm: str, timeout: int, aggressive: bool = False) -> str:
    best_text = ""
    best_score = float("-inf")
    for variant in _preprocess_variants(image, aggressive=aggressive):
        try:
            text = pytesseract.image_to_string(
                variant,
                lang=OCR_LANG,
                config=_ocr_config(psm),
                timeout=timeout,
            )
        except RuntimeError:
            text = "[OCR_TIMEOUT]"
        except Exception as error:
            # If the Vietnamese tessdata pack is missing on a deployed machine,
            # retry English so the pipeline still returns safe metadata instead
            # of crashing.  Quality guards later prevent bad OCR from leaking.
            if "vie" in OCR_LANG.lower():
                try:
                    text = pytesseract.image_to_string(
                        variant,
                        lang="eng",
                        config=_ocr_config(psm),
                        timeout=timeout,
                    )
                except Exception:
                    text = f"[OCR_ERROR: {type(error).__name__}]"
            else:
                text = f"[OCR_ERROR: {type(error).__name__}]"
        score = _score_ocr_text(text)
        if score > best_score:
            best_text = text
            best_score = score
    return best_text.strip()


def _page_to_image(page, zoom: float) -> Image.Image:
    matrix = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=matrix, alpha=False)
    return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)


def ocr_pdf_page(page, page_index: int) -> str:
    image = _page_to_image(page, OCR_RENDER_ZOOM)
    text = _image_to_string_best(image, psm=OCR_TESSERACT_PSM, timeout=OCR_TIMEOUT_SEC, aggressive=False)
    return f"\n--- PAGE {page_index} OCR ---\n{text.strip()}"


def ocr_pdf_targeted_region(
    page,
    page_index: int,
    label: str,
    crop_box_ratio: Tuple[float, float, float, float],
    psm: str,
) -> str:
    image = _page_to_image(page, OCR_TARGETED_ZOOM)
    width, height = image.size

    left, top, right, bottom = crop_box_ratio
    crop = image.crop((
        max(0, int(width * left)),
        max(0, int(height * top)),
        min(width, int(width * right)),
        min(height, int(height * bottom)),
    ))
    text = _image_to_string_best(crop, psm=psm, timeout=OCR_TARGETED_TIMEOUT_SEC, aggressive=True)
    return f"\n--- PAGE {page_index} {label} ---\n{text.strip()}"


def build_targeted_ocr_text(page, page_index: int, page_count: int, page_zero_index: int, should_ocr: bool) -> str:
    if not OCR_TARGETED_REGIONS or not (should_ocr or OCR_TARGETED_ALWAYS):
        return ""

    parts: List[str] = []
    if page_zero_index == 0:
        # Separate zones reduce stamp/seal pollution in long fields.
        parts.append(ocr_pdf_targeted_region(page, page_index, "DOC_CODE_OCR", (0.00, 0.05, 0.50, 0.22), "6"))
        parts.append(ocr_pdf_targeted_region(page, page_index, "DATE_LINE_OCR", (0.42, 0.07, 0.98, 0.20), "6"))
        parts.append(ocr_pdf_targeted_region(page, page_index, "HEADER_OCR", (0.00, 0.04, 1.00, 0.27), "6"))
        parts.append(ocr_pdf_targeted_region(page, page_index, "TITLE_OCR", (0.06, 0.14, 0.94, 0.40), "6"))
        # Công văn/VPCP files usually put V/v and Kính gửi in the left/middle block.
        parts.append(ocr_pdf_targeted_region(page, page_index, "VV_BLOCK_OCR", (0.03, 0.10, 0.91, 0.46), "6"))
        # Article 1 often repeats/expands the operative subject for Quyết định/Nghị quyết.
        parts.append(ocr_pdf_targeted_region(page, page_index, "ARTICLE1_OCR", (0.08, 0.30, 0.95, 0.82), "6"))

    if page_zero_index == page_count - 1:
        # Signature, signer name/title and Lưu code.
        parts.append(ocr_pdf_targeted_region(page, page_index, "SIGNATURE_FOOTER_OCR", (0.02, 0.52, 0.99, 0.96), "6"))

    return "\n".join(part for part in parts if part.strip())


PdfReadMode = Literal["first_page", "first_and_last_page", "full_pdf"]


def build_pdf_page_indices(page_count: int, pdf_read_mode: PdfReadMode) -> List[int]:
    if pdf_read_mode == "full_pdf":
        if page_count <= 2:
            return list(range(page_count))
        return [0, page_count - 1] + list(range(1, page_count - 1))
    if pdf_read_mode == "first_and_last_page":
        return [0] if page_count == 1 else [0, page_count - 1]
    return [0]


def extract_text_from_pdf(file_path: Path, pdf_read_mode: PdfReadMode = "first_and_last_page") -> Tuple[str, int]:
    page_records: List[Tuple[int, str]] = []
    ocr_used_pages = 0

    with fitz.open(file_path) as pdf:
        page_count = len(pdf)
        if page_count == 0:
            return "", 0

        page_indices = build_pdf_page_indices(page_count, pdf_read_mode)
        for page_zero_index in page_indices:
            page_index = page_zero_index + 1
            page = pdf[page_zero_index]
            text_layer = page.get_text("text").strip()

            should_ocr = is_text_layer_noise(text_layer)
            if should_ocr and OCR_MAX_PAGES_PER_PDF > 0 and ocr_used_pages >= OCR_MAX_PAGES_PER_PDF:
                should_ocr = False

            ocr_text = ""
            targeted_ocr_text = ""
            if should_ocr or OCR_TARGETED_ALWAYS:
                targeted_ocr_text = build_targeted_ocr_text(page, page_index, page_count, page_zero_index, should_ocr)
                # Full-page OCR is slow on scanned PDFs.  For metadata extraction,
                # targeted regions are normally enough.  Enable OCR_FULL_PAGE=1
                # only when you need body text from every page.
                if should_ocr and OCR_FULL_PAGE:
                    ocr_text = ocr_pdf_page(page, page_index)
                if should_ocr:
                    ocr_used_pages += 1

            page_text = f"""
--- PAGE {page_index} / {page_count} ---

[TEXT_LAYER]
{text_layer}

[OCR_TEXT]
{ocr_text}
{targeted_ocr_text}
""".strip()
            page_records.append((page_zero_index, page_text))

    pages = [page_text for _page_zero_index, page_text in sorted(page_records, key=lambda item: item[0])]
    return "\n\n".join(pages).strip(), page_count


def extract_text_from_docx(file_path: Path) -> Tuple[str, Optional[int]]:
    document = Document(file_path)
    parts: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"\n--- TABLE {table_index} ---")
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                cells.append(cell_text)
            row_text = " | ".join(cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts).strip(), None


def extract_text_from_txt(file_path: Path) -> Tuple[str, Optional[int]]:
    return file_path.read_text(encoding="utf-8", errors="ignore").strip(), None


def extract_text_from_image(file_path: Path) -> Tuple[str, Optional[int]]:
    with Image.open(file_path) as image:
        text = _image_to_string_best(image.convert("RGB"), psm=OCR_TESSERACT_PSM, timeout=OCR_TIMEOUT_SEC, aggressive=True)
    return text.strip(), 1


def extract_text(file_path: Path, pdf_read_mode: PdfReadMode = "first_and_last_page") -> Tuple[str, Optional[int]]:
    ext = file_path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        return extract_text_from_pdf(file_path, pdf_read_mode=pdf_read_mode)
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    if ext == ".txt":
        return extract_text_from_txt(file_path)
    if ext in {".png", ".jpg", ".jpeg"}:
        return extract_text_from_image(file_path)

    raise ValueError(f"Unsupported file type: {ext}")
